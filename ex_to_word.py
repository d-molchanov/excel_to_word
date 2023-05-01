from os import walk
from os.path import abspath, relpath, splitext, join, basename

from time import perf_counter

import openpyxl
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL

from directive import Directive

def timeit(operation):
    def decorator(method):
        # def wrapped(filename):
        def wrapped(*args, **kwargs):
            base_name = basename(kwargs['_path'])
            print(f'{operation} <{base_name}>...\r')
            time_start = perf_counter()
            result = method(*args, **kwargs)
            time_finish = (perf_counter() - time_start)*1e3
            print(f'{operation} <{base_name}> done in {time_finish:.3f} ms.')
            return result
        return wrapped
    return decorator

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_xlsx_file_list(target_dir):
    result = []
    for root, dirs, files in walk(target_dir):
        for f in files:
            ap_file = abspath(join(root, f))
            filename, ext = splitext(ap_file)
            if ext == '.xlsx':
                result.append(ap_file)
    return result

@timeit('Reading')
def read_xlsx(_path):
    wb = openpyxl.load_workbook(_path)
    sheet = wb.active
    rows = sheet.rows
    data = [[cell.value for cell in row] for row in sheet.rows]
    return [row for row in data if row != [None for el in row]]

def extract_columns(data, columns):
    return [[row[el] for el in columns] for row in data] 

def change_ext(filename, new_ext):
    name, ext = splitext(filename)
    return f'{name}.{new_ext}'


def convert_data_to_str(data, formatting):
    return [
        [f.format(r).replace('.', ',') if type(r) != str else r for r, f in 
        zip(row, formatting)] for row in data
    ]

def add_table_new(document, data, koord_zone):
    table_section = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = table_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    h = len(data)
    w = len(data[0])

    table = document.add_table(rows=h+3, cols=w, style='Table Grid')
    table.style.font.name = 'Times New Roman'
    table.style.font.size = Pt(12)
    table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells = table._cells

    cells[0].merge(cells[3])
    cells[1].merge(cells[2])

    index = koord_zone.find(')', None)
    if index:
        koord_zone = koord_zone[:index+1]

    head = [
        'Обозначение характерных точек границ',
        koord_zone,
        'X',
        'Y',
        '(1)',
        '(2)',
        '(3)'
    ]

    head_cells = [0, 1, 4, 5, 6, 7, 8]

    for t, c in zip(head, [cells[el] for el in head_cells]):
        c.text = t
        c.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].style.font.name = 'Times New Roman'
        c.paragraphs[0].style.font.size = Pt(11)
        c.paragraphs[0].runs[0].bold = True


    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].height = Cm(1.96)

    # table.allow_autofit = False
    table.columns[0].width = Cm(4)
    for i, c in enumerate(cells[9:]):
        c.text = data[i//w][i%w]
        # p = c.paragraphs[0]
        # p.style.font.name = 'Times New Roman'
        # p.style.font.size = Pt(12)
        # p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_section = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = footer_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'1')

    p = document.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

    return document

def scan_directory(target_dir, filenames):
    result = dict()
    for root, dirs, files in walk(target_dir):
        found_files = [f for f in files if f in filenames]
        found_files.sort()
        if found_files == filenames:
            result[f'{abspath(root)}'] = found_files
    # print(result)
    return result

@timeit('Writing')
def write_txtfile(data, sep, _path):
    try:
        with open(_path, 'w') as f:
            for row in data:
                f.write(f"{sep.join(row)}\n")
    except IOError:
        print(f'I/O error with <{_path}>.')

@timeit('Reading')
def read_textfile(_path):
    try:
        with open(_path, 'r', encoding='utf-8') as f:
            return [line.rstrip() for line in f.readlines()]
    except IOError:
        print(f'I/O error with <{_path}>.')
        return None 

def set_page_properties(document):
    header = document.sections[0]
    header.page_width = Cm(21)
    header.page_height = Cm(29.7)
    header.left_margin = Cm(3)
    header.right_margin = Cm(1.5)
    header.top_margin = Cm(2)
    header.bottom_margin = Cm(2)

def set_directive_styles(document):
    main_style = document.styles.add_style(
        'Directive Text', WD_STYLE_TYPE.PARAGRAPH)
    main_style.next_paragraph_style = main_style
    main_style.font.size = Pt(14)
    main_style.font.name = 'Times New Roman'
    p_f = main_style.paragraph_format
    p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_f.first_line_indent = Cm(1.25)
    p_f.line_spacing = 1.15
    p_f.space_before = Pt(0)
    p_f.space_after = Pt(0)

    title_style = document.styles.add_style(
        'Directive Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.base_style = main_style
    title_style.next_paragraph_style = main_style
    title_style.font.bold = True
    p_f = title_style.paragraph_format
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.line_spacing = 1.15
    p_f.first_line_indent = Cm(0)
    p_f.space_before = Pt(238)
    p_f.space_after = Pt(42)

    position_style = document.styles.add_style(
        'Directive Position', WD_STYLE_TYPE.PARAGRAPH)
    position_style.base_style = main_style
    p_f = position_style.paragraph_format
    p_f.first_line_indent = Cm(0)
    p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT


    name_style = document.styles.add_style(
        'Directive Name', WD_STYLE_TYPE.PARAGRAPH)
    name_style.base_style = main_style
    p_f = name_style.paragraph_format
    p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return [title_style, main_style, position_style, name_style]

def create_position_table(document, data, styles):
    table = document.add_table(rows=1, cols=2)
    cells = table._cells
    for d, s, c in zip(data, styles, cells):
        c.text = d
        c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        c.paragraphs[0].style = s

# def find_all_substring_indices(string, substring):
#     i = 0
#     result = []
#     while True:
#         i = string.find(substring, i)
#         if i == -1:
#             break
#         result.append(i)
#         i+=len(substring)
#         result.append(i)
#     return result

# # !Check this method
# def split_string(string, substrings):
#     indices = [0]
#     for s in substrings:
#         indices += find_all_substring_indices(string, s)
#     indices.sort()
#     if indices[-1] != len(string) - 1:
#         # indices.append(len(string) - 1)
#         indices.append(len(string))
#     result = []
#     for i, j in zip(indices[1:], indices[:-1]):
#         result.append(string[j:i])
#     return result

def enumerate_part_of_directive(part_number, part_of_directive):
    result = [f'{part_number}. {part_of_directive[0]}']
    if len(part_of_directive) > 1:
        for i, p in enumerate(part_of_directive[1:], 1):
            result.append(f'{part_number}.{i}. {p}')
    return result

def create_directive_text(directive_template, appendixes_2_and_3_are_equal):
    result = directive_template[:2]
    part_1 = directive_template[2:3]
    part_2 = directive_template[3:4]
    if appendixes_2_and_3_are_equal:
        part_2.append(directive_template[6])
    else:
        part_2 += directive_template[4:6]
    part_2 += directive_template[7:8]
    part_3 = directive_template[8:9]
    part_4 = directive_template[9:17]
    part_5 = directive_template[17:18]
    parts = [part_1, part_2, part_3, part_4, part_5]
    for i, p in enumerate(parts, 1):
        result += enumerate_part_of_directive(i, p)
    result += directive_template[-4:]
    return result

def create_directive_new(directive_template, substitution):
    document = Document()
    set_page_properties(document)
    
    [
        title_style, main_style,
        position_style, name_style
    ] = set_directive_styles(document)

    water_object = substitution[0]
    district_name = substitution[1]
    water_object_length = substitution[2]
    water_protection_zone = substitution[3]
    protected_shoreline_belt = substitution[4]

    subst_templ = {
        '{WO}': water_object,
        '{DN}': district_name,
        '{WOL}': water_object_length,
        '{WPZ}': water_protection_zone,
        '{PSB}': protected_shoreline_belt
    }

    if directive_template:
        for line in directive_template[:-4]:
            p = document.add_paragraph(style=main_style)
            subline = split_string(line, subst_templ)
            for s in subline:
                 if s in subst_templ:
                    text = subst_templ.get(s)
                    if text:
                        p.add_run(text)
                    else:
                        r = p.add_run(s)
                        r.font.highlight_color = WD_COLOR_INDEX.RED
                 else:
                    p.add_run(s.format(NBS=chr(160)))

        document.paragraphs[0].style = title_style
        for i in range(3):
            document.add_paragraph(style=main_style)
        
        position = '\n'.join(directive_template[-4:-1])
        name = directive_template[-1]

        create_position_table(
            document, [position, name],
            [position_style, name_style])
        document.add_page_break()

        return document
        
    else:
        print(f'<{directive_template}> is empty.')
        return None

#!Rewrite with list comprehension
def create_document_framework(template_data, indices, separators):
    result = []
    for i, j, s in zip(indices[:-1], indices[1:], separators):
        result.append(s.join(template_data[i:j]))
    return result

def create_appendix_content(framework, appendix_number):
    indices = []
    if appendix_number == '1':
        indices += [2, 6]
    elif appendix_number == '2':
        indices += [3, 7]
    elif appendix_number == '3':
        indices += [4, 8]
    else:
        indices += [5, 9]
        appendix_number = '2'
    result = [
        framework[0].format(AN=appendix_number),
        framework[1]
    ]
    for i in indices:
        result.append(framework[i])
    return result

def set_appendix_styles(document):
    text_style = document.styles.add_style(
        'Appendix Text', WD_STYLE_TYPE.PARAGRAPH)
    text_style.font.size = Pt(14)
    text_style.font.name = 'Times New Roman'
    text_style.next_paragraph_style = text_style
    p_f = text_style.paragraph_format
    p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_f.first_line_indent = Cm(1.25)
    p_f.line_spacing = 1.15
    p_f.space_before = Pt(0)
    p_f.space_after = Pt(14)

    number_style = document.styles.add_style(
        'Appendix Document Number', WD_STYLE_TYPE.PARAGRAPH)
    number_style.base_style = text_style
    number_style.font.size = Pt(13)
    p_f = number_style.paragraph_format
    p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_f.first_line_indent = Cm(0)
    p_f.left_indent = Cm(11)
    
    appendix_style = document.styles.add_style(
        'Appendix Title', WD_STYLE_TYPE.PARAGRAPH)
    appendix_style.base_style = number_style
    p_f = appendix_style.paragraph_format
    # p_f.left_indent = Cm(11)
    p_f.space_after = Pt(0)
    
    title_style = document.styles.add_style(
        'Appendix Document Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.base_style = text_style
    title_style.font.bold = True
    p_f = title_style.paragraph_format
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.first_line_indent = Cm(0)
    
    return [appendix_style, number_style, title_style, text_style]

def add_appendix(document, content):
    appendix_styles = [
        document.styles['Appendix Title'],
        document.styles['Appendix Document Number'],
        document.styles['Appendix Document Title'],
        document.styles['Appendix Text']
    ]
    document.add_section()
    for c, s in zip(content, appendix_styles):
        document.add_paragraph(c, style=s)
    return document


def write_docx_file(document, output_file):
    try:
        document.save(output_file)
    except PermissionError:
        print(f'<{output_file}> is busy - permission denied.')


def process_directory(target_dir, filenames):
    dir_content = scan_directory(target_dir, filenames)
    appendix_content = read_textfile('appendix_template.txt')
    appendix_framework = create_document_framework(
        appendix_content,
        [0, 5, 6, 9, 12, 15, 18, 19, 20, 21, 22],
        ['\n', '', '\n', '\n', '\n', '\n', '', '', '', '']
    )

    for k, v in dir_content.items():
        # print(f'{k}:\t{v}')
        print(f'Start processing <{k}>:')
        t_s = perf_counter()

        print('Reading <content.txt>...', end='\r')
        time_start = perf_counter()
        content_txt = read_textfile(join(k, 'content.txt'))
        time_finish = round((perf_counter() - time_start)*1e3, 3)
        print(f'Reading <content.txt> done in {time_finish} ms.')
        
        formatted_content = [None for i in range(5)]
        if content_txt:
            for i, c in enumerate(content_txt):
                formatted_content[i] = c.format(NBS=chr(160))
        water_object_name = formatted_content[0]        
        # print(formatted_content)
        xlsx_files = [f for f in v if splitext(f)[1] == '.xlsx']
        xlsx_files.sort()
        appendix_document = Document()
        xlsx_data = []
        koord_zone = []
        appendix_numbers = []
        for i, f in enumerate(xlsx_files, 1):
            # appendix_numbers.append(f[-6])
            appendix_numbers.append(str(i))
            
            print(f'Reading <{f}>...', end='\r')
            time_start = perf_counter()
            data = read_xlsx(join(k, f))
            time_finish = round((perf_counter() - time_start)*1e3, 3)
            print(f'Reading <{f}> done in {time_finish} ms ({len(data)} rows).')
            
            xlsx_data.append(data)
            koord_zone.append(data[2][1])
        partial_data = [extract_columns(el[5:], [0, 1, 2]) for el in xlsx_data]
        formatting = ['{:.0f}', '{:.2f}', '{:.2f}']
        str_data = [convert_data_to_str(el, formatting) for el in partial_data]

        for f, d in zip(xlsx_files, str_data):
            new_filename = change_ext(f, 'txt')
            time_start = perf_counter()
            print(f'Writing <{new_filename}>...', end='\r')
            time_start = perf_counter()
            write_txtfile(extract_columns(d, [1, 2]), join(k, new_filename), '\t')
            time_finish = round((perf_counter() - time_start)*1e3, 3)
            print(f'Writing <{new_filename}> done in {time_finish} ms ({len(d)} rows).')


        appendixes_2_and_3_are_equal = False
        if partial_data[1] == partial_data[2]:
            appendix_numbers = ['1', '23']
            appendixes_2_and_3_are_equal = True
        
        print('Reading <directive_template.txt>...', end='\r')
        time_start = perf_counter()
        directive_template = read_textfile('directive_template.txt')
        time_finish = round((perf_counter() - time_start)*1e3, 3)
        print(f'Reading <directive_template.txt> done in {time_finish} ms.')
        
        directive_text = create_directive_text(directive_template, appendixes_2_and_3_are_equal)
        
        print('Creating directive...', end='\r')
        time_start = perf_counter()
        document = create_directive_new(directive_text, formatted_content)
        time_finish = round((perf_counter() - time_start)*1e3, 3)
        print(f'Reading directive done in {time_finish} ms.')

        paragraph = document.sections[0].footer.paragraphs[0]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.size = Pt(12)
        paragraph.style.font.name = 'Times New Roman'
        add_page_number(paragraph.add_run())

        set_appendix_styles(document)
        for d, k_z, n in zip(str_data, koord_zone, appendix_numbers):
            a_c = create_appendix_content(appendix_framework, water_object_name, n)
            print(f'Creating appendix {n}...', end='\r')
            time_start = perf_counter()
            document = add_appendix(document, a_c)
            document = add_table_new(document, d, k_z)
            time_finish = round((perf_counter() - time_start)*1e3, 3)
            print(f'Creating appendix {n} done in {time_finish} ms.')

        for s in document.sections:
            sectPr = s._sectPr
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(ns.qn('w:start'), "1")
            sectPr.append(pgNumType)

        output_file = 'All_Appendix.docx'
        print(f'Writing <{output_file}>...', end='\r')
        time_start = perf_counter()
        write_docx_file(document, join(k, output_file))
        time_finish = round((perf_counter() - time_start)*1e3, 3)
        print(f'Writing <{output_file}> done in {time_finish} ms.')

        t_f = round((perf_counter() - t_s)*1e3, 3)
        print(f'Processing <{k}> done in {t_f} ms.')

    processed_dirs = list(dir_content.keys())
    processed_dirs.sort()
    print('Processed directories:', *processed_dirs, sep='\n')

def process_waterbody(target_dir, filenames, directive_template, appendix_framework):
    print(f'Start processing <.{relpath(target_dir, ".")}>:')
    t_s = perf_counter()

    content_txt = read_textfile(_path=join(target_dir, 'content.txt'))
    subst_keys = ['WBN', 'DN', 'WBL', 'WPZ', 'PSB']
    substitution = {k:f'{{{k}}}' for k in subst_keys}
    if content_txt:
        # for k, c in zip(subst_keys[1:], content_txt[1:]):
        for k, c in zip(subst_keys, content_txt):
            substitution[k] = c

    waterbody_name = substitution['WBN']

    xlsx_files = [f for f in filenames if splitext(f)[1] == '.xlsx']
    xlsx_files.sort()
    appendix_document = Document()
    xlsx_data = []
    coordinates_title = []
    appendix_numbers = []
    for i, f in enumerate(xlsx_files, 1):
        appendix_numbers.append(str(i))
        data = read_xlsx(_path=join(target_dir, f))
        xlsx_data.append(data)
        coordinates_title.append(data[2][1])
    coordinates = [extract_columns(el[5:], [0, 1, 2]) for el in xlsx_data]
    formatting = ['{:.0f}', '{:.2f}', '{:.2f}']
    str_coordinates = [convert_data_to_str(el, formatting) for el in coordinates]

    for f, d in zip(xlsx_files, str_coordinates):
        new_filename = change_ext(f, 'txt')
        write_txtfile(extract_columns(d, [1, 2]), _path=join(target_dir, new_filename), sep='\t')


    appendixes_2_and_3_are_equal = False
    if coordinates[1] == coordinates[2]:
        appendix_numbers = ['1', '23']
        appendixes_2_and_3_are_equal = True
    
    directive_text = create_directive_text(directive_template, appendixes_2_and_3_are_equal)
    directive = Directive()
    doc = directive._create_directive(directive_text, substitution)
    directive._set_page_properties(doc)
    
#     print('Creating directive...', end='\r')
#     time_start = perf_counter()
#     document = create_directive_new(directive_text, formatted_content)
#     time_finish = round((perf_counter() - time_start)*1e3, 3)
#     print(f'Reading directive done in {time_finish} ms.')


#     set_appendix_styles(document)
    for c, c_t, n in zip(str_coordinates, coordinates_title, appendix_numbers):
        a_c = create_appendix_content(appendix_framework, n)
        doc = directive._add_appendix(doc, a_c, substitution)
        doc = directive._add_table(doc, c, c_t)
#         print(f'Creating appendix {n}...', end='\r')
#         time_start = perf_counter()
#         document = add_appendix(document, a_c)
#         document = add_table_new(document, d, k_z)
#         time_finish = round((perf_counter() - time_start)*1e3, 3)
#         print(f'Creating appendix {n} done in {time_finish} ms.')
    directive._create_right_numeration(doc)
#     for s in document.sections:
#         sectPr = s._sectPr
#         pgNumType = OxmlElement('w:pgNumType')
#         pgNumType.set(ns.qn('w:start'), "1")
#         sectPr.append(pgNumType)

#     output_file = 'All_Appendix.docx'
#     print(f'Writing <{output_file}>...', end='\r')
#     time_start = perf_counter()
#     write_docx_file(document, join(k, output_file))
#     time_finish = round((perf_counter() - time_start)*1e3, 3)
#     print(f'Writing <{output_file}> done in {time_finish} ms.')

#     t_f = round((perf_counter() - t_s)*1e3, 3)
#     print(f'Processing <{k}> done in {t_f} ms.')
    doc.save('Directive_Templ.docx')

# processed_dirs = list(dir_content.keys())
# processed_dirs.sort()
# print('Processed directories:', *processed_dirs, sep='\n')

def process_directory_new(target_dir, filenames):
    dir_content = scan_directory(target_dir, filenames)

    directive_template = read_textfile(_path='directive_template.txt')
    appendix_content = read_textfile(_path='appendix_template.txt')
    
    appendix_framework = create_document_framework(
        appendix_content,
        [0, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14],
        ['\n', '', '', '', '', '', '', '', '', '']
    )

    for directory, files in dir_content.items():
        process_waterbody(directory, files, directive_template, appendix_framework)

if __name__ == '__main__':
    # target_dir = './data/26_река_Нахавня_(Одинцовские г.о.)'
    # target_dir = './data/10_ручей_без_названия_(г.о. Егорьевск)'
    # target_dir = './data/12_река_Плесенка_(Наро-Фоминский г.о.)'
    # target_dir = './data/(2023_04_02)/6_река_Вьюница_(г.о. Шатура)'
    # target_dir = './data/(2023_04_02)/13_река_Шатуха_(Наро-Фоминский г.о., Рузский г.о.)'
    # target_dir = './data/(2023_04_02)/16_река_Сумерь_(г.о. Пушкинский, Сергиево-Посадский г.о.)'
    # target_dir = './data/(2023_04_02)/26_река_Нахавня_(Одинцовские г.о.)'
    # target_dir = './data/(2023_04_02)/33_река_Малые_Вяземы_(Одинцовский г.о.)'
    # target_dir = './data/(2023_04_02)/34_ручей_без_названия_(г.о. Домодедово)'
    # target_dir = './data/(2023_04_02)/37_ручей_без_названия_(г.о. Домодедово)'
    # target_dir = './data/(2023_04_02)/42_река_Беляна_(Одинцовский г.о., г.о. Истра)'
    # target_dir = './data/(2023_04_02)/44_река_Жданка_(Раменский г.о., г.о. Домодедово)'
    # target_dir = './data/(2023_04_02)/46_река_Лубянка_(г.о. Ступино)'
    # target_dir = './data/(2023_04_02)/49_река_Камариха_(г.о. Пушкинский, Дмитровский г.о.)'
    # target_dir = './data/(2023_04_02)/50_река_Вырка_(Орехово-Зуевский г.о.)'

    # target_dir ='./data/Проекты распоряжений'
    target_dir ='./data/Проекты распоряжений/Лотошинский район/299 река Черная test'
    # target_dir ='./data/Проекты распоряжений/Лотошинский район/299 река Черная'
    # target_dir ='./data/Проекты распоряжений/Лотошинский район/865 река Безымянная'

    # filenames = [
    #     'Приложение 1.xlsx',
    #     'Приложение 2.xlsx',
    #     'Приложение 3.xlsx',
    #     'content.txt'
    # ]

    filenames = [
        'каталог координат БЛ.xlsx',
        'каталог координат ВОЗ.xlsx',
        'каталог координат ПЗП.xlsx',
        'content.txt'
    ]
    filenames.sort()

    time_start = perf_counter()
    print(f'Processing <{target_dir}>...')
    process_directory_new(target_dir, filenames)
    time_finish = round((perf_counter() - time_start)*1e3, 3)
    print(f'Processing <{target_dir}> done in {time_finish} ms.')
    