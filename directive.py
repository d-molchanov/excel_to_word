from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL

class Directive:

    

    def _create_element(self, name):
        return OxmlElement(name)

    def _create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)


    def _add_page_number(self, run):
        fldChar1 = self._create_element('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self._create_element('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def _set_page_properties(self, document):
        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def _set_directive_styles(self, document):
        main_style = document.styles.add_style(
            'Directive Text', WD_STYLE_TYPE.PARAGRAPH)
        main_style.next_paragraph_style = main_style
        main_style.font.size = Pt(14)
        main_style.font.name = 'Times New Roman'
        p_f = main_style.paragraph_format
        p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_f.first_line_indent = Cm(1.25)
        p_f.line_spacing = 1.15
        p_f.space_before = Pt(0)
        p_f.space_after = Pt(0)

        title_style = document.styles.add_style(
            'Directive Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = main_style
        title_style.next_paragraph_style = main_style
        title_style.font.bold = True
        p_f = title_style.paragraph_format
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.line_spacing = 1.15
        p_f.first_line_indent = Cm(0)
        p_f.space_before = Pt(238)
        p_f.space_after = Pt(42)

        position_style = document.styles.add_style(
            'Directive Position', WD_STYLE_TYPE.PARAGRAPH)
        position_style.base_style = main_style
        p_f = position_style.paragraph_format
        p_f.first_line_indent = Cm(0)
        p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT


        name_style = document.styles.add_style(
            'Directive Name', WD_STYLE_TYPE.PARAGRAPH)
        name_style.base_style = main_style
        p_f = name_style.paragraph_format
        p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        return [title_style, main_style, position_style, name_style]

    def _set_appendix_styles(self, document):
        text_style = document.styles.add_style(
            'Appendix Text', WD_STYLE_TYPE.PARAGRAPH)
        text_style.font.size = Pt(14)
        text_style.font.name = 'Times New Roman'
        text_style.next_paragraph_style = text_style
        p_f = text_style.paragraph_format
        p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_f.first_line_indent = Cm(1.25)
        p_f.line_spacing = 1.15
        p_f.space_before = Pt(0)
        p_f.space_after = Pt(14)

        number_style = document.styles.add_style(
            'Appendix Document Number', WD_STYLE_TYPE.PARAGRAPH)
        number_style.base_style = text_style
        number_style.font.size = Pt(13)
        p_f = number_style.paragraph_format
        p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_f.first_line_indent = Cm(0)
        p_f.left_indent = Cm(11)
        
        appendix_style = document.styles.add_style(
            'Appendix Title', WD_STYLE_TYPE.PARAGRAPH)
        appendix_style.base_style = number_style
        p_f = appendix_style.paragraph_format
        # p_f.left_indent = Cm(11)
        p_f.space_after = Pt(0)
        
        title_style = document.styles.add_style(
            'Appendix Document Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = text_style
        title_style.font.bold = True
        p_f = title_style.paragraph_format
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.first_line_indent = Cm(0)
        
        return [appendix_style, number_style, title_style, text_style]

    def _create_position_table(self, document, data, styles):
        table = document.add_table(rows=1, cols=2)
        cells = table._cells
        for d, s, c in zip(data, styles, cells):
            c.text = d
            c.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            c.paragraphs[0].style = s

    def _create_mask(self, text_list, substitution):
        flags = []
        text = []
        for el in text_list:
            subs = substitution.get(el, None)
            if subs == None:
                flags.append(0)
            elif subs == el:
                flags.append(2)
            else:
                flags.append(1)
                el = subs
            text.append(el)
        return {
            'flags': flags,
            'text': text
        }

    def _apply_mask(self, paragraph, text_list, flags):
        for t, f in zip(text_list, flags):
            r = paragraph.add_run(t)
            if f == 1:
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif f == 2:
                r.font.highlight_color = WD_COLOR_INDEX.RED

    def __init__(self):
        self._document = Document()
        self._set_page_properties(self._document)
        self._set_directive_styles(self._document)

    def _add_appendix(self, document, content, substitution):
        appendix_styles = [
            document.styles['Appendix Title'],
            document.styles['Appendix Document Number'],
            document.styles['Appendix Document Title'],
            document.styles['Appendix Text']
        ]
        document.add_section()

        content_mask = self._create_template_mask(
            content, substitution, '#')
        text = content_mask['text']
        flags = content_mask['flags']

        for line, flag_line, s in zip(text, flags, appendix_styles):
            p = document.add_paragraph(style=s)
            self._apply_mask(p, line, flag_line)


        return document

    def _create_template_mask(self, template, _substitution, _sep):
        split_substitution = {
            k: f'{_sep}{{{k}}}{_sep}' for k in _substitution.keys()}
        text = []
        flags = []
        substitution = {
            f'{{{k}}}': v.replace('{NBS}', chr(160)) for 
            k, v in _substitution.items()}
        for line in template:
            line = line.replace('{NBS}', chr(160))
            text_list = [
                el for el in line.format(
                    **split_substitution).split(_sep) if el]
            mask = self._create_mask(text_list, substitution)
            text.append(mask['text'])
            flags.append(mask['flags'])

        return {
            'text': text,
            'flags': flags
        }

    def _create_directive(self, directive_template, substitution):
        directive_mask = self._create_template_mask(
            directive_template, substitution, '#')

        text = directive_mask['text']
        flags = directive_mask['flags']


        document = Document()
        self._set_page_properties(document)
        
        [
            title_style, main_style,
            position_style, name_style
        ] = self._set_directive_styles(document)

        if text:
            for line, flag_line in zip(text[:-4], flags[:-4]):
                p = document.add_paragraph(style=main_style)
                self._apply_mask(p, line, flag_line)

        document.paragraphs[0].style = title_style
        
        for i in range(3):
            document.add_paragraph(style=main_style)
            
        position = '\n'.join([el[0] for el in text[-4:-1]])
        name = text[-1][0]

        self._create_position_table(
            document, [position, name],
            [position_style, name_style])
        document.add_page_break()

        paragraph = document.sections[0].header.paragraphs[0]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.size = Pt(12)
        paragraph.style.font.name = 'Times New Roman'
        self._add_page_number(paragraph.add_run())

        self._set_appendix_styles(document)
        return document
            
    def _create_right_numeration(self, document):
        for s in document.sections:
            sectPr = s._sectPr
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(ns.qn('w:start'), "1")
            sectPr.append(pgNumType)
            s.different_first_page_header_footer = True

    def _add_table(self, document, data, coordinates_title):
        table_section = document.add_section(WD_SECTION.CONTINUOUS)
        sectPr = table_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')

        h = len(data)
        w = len(data[0])

        table = document.add_table(rows=h+3, cols=w, style='Table Grid')
        table.style.font.name = 'Times New Roman'
        table.style.font.size = Pt(12)
        table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells = table._cells

        cells[0].merge(cells[3])
        cells[1].merge(cells[2])

        index = coordinates_title.find(')', None)
        if index:
            coordinates_title = coordinates_title[:index+1]

        head = [
            'Обозначение характерных точек границ',
            coordinates_title,
            'X',
            'Y',
            '(1)',
            '(2)',
            '(3)'
        ]

        head_cells = [0, 1, 4, 5, 6, 7, 8]

        for t, c in zip(head, [cells[el] for el in head_cells]):
            c.text = t
            c.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].style.font.name = 'Times New Roman'
            c.paragraphs[0].style.font.size = Pt(11)
            c.paragraphs[0].runs[0].bold = True


        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.rows[0].height = Cm(1.96)

        table.columns[0].width = Cm(4)
        for i, c in enumerate(cells[9:]):
            c.text = data[i//w][i%w]

        footer_section = document.add_section(WD_SECTION.CONTINUOUS)
        sectPr = footer_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'1')

        p = document.add_paragraph()
        r = p.add_run()
        r.add_break(WD_BREAK.PAGE)

        return document         

